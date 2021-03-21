import docx
import re
import math
import matplotlib.pyplot as plt
from urllib.request import urlretrieve
from PIL import Image
from resizeimage import resizeimage
import logging
from os import path, getcwd

# create the the log file path that works on all OS
file_path = path.join(getcwd(), 'logging.txt')
# Create logger
LOG_FORMAT = '%(levelname)s %(asctime)s %(filename)s - %(message)s'
logging.basicConfig(filename=file_path, level=logging.DEBUG, format=LOG_FORMAT)
# root logger
logger = logging.getLogger()
File_NAME = "book.txt"


# Task 2
def read_authors(filename):
    """read the author his information and the book title"""
    dict_file_content = dict()
    file_content = None
    # reading the file in a safe way
    try:
        with open(filename, encoding='utf-8') as infile:
            file_content = infile.read()
            title_pattern = re.compile(r"(Title:)(.*)")
            author_pattern = re.compile(r"(Author:)(.*)")
            title = title_pattern.search(file_content)
            author = author_pattern.search(file_content)
    except OSError:
        print("Something went wrong")
        exit()
    finally:
        book_title = str()
        book_author = str()
        if(title is None):
            book_title = "Unknown"
        else:
            book_title = title.group(2)
        if(author is None):
            book_author = "Unknown"
        else:
            book_author = author.group(2)
        return (book_title, book_author)


# Task 3
def get_plot_data(filename):
    """read the first chapter of the paragrap and return a list containing
    number of words"""
    dict_file_content = dict()
    file_content = []
    # reading the file in a safe way
    try:
        with open(filename, encoding='utf-8') as infile:
            file_content = infile.read()
            # make the . match even new Lines
            chap_pattern = re.compile(r"(CHAPTER I)(.*?)(CHAPTER II)",
                                      re.DOTALL)
            # search pattern
            first_chap = (chap_pattern.search(file_content))
    except OSError:
        print("Something went wrong")
        exit()
    finally:
        # dived the number paragraph
        chapters = str(first_chap.group(2)).split("\n\n")

        def round_up(x):
            return int(math.ceil(x / 10.0)) * 10
        # avoid duplicate  I could use set but not asked
        s = []
        for chap in chapters:
            s.append(round_up(len(chap.split())))
        return sorted(s)


# Task 3b
def plot_linear_chart(data):
    """Getting data from get_plot_data and plot a figure based on those data"""
    dict_data = dict(
        sorted(dict((i, data.count(i)) for i in set(data)).items()))
    keys_list = list(dict_data.keys())
    values_list = list(dict_data.values())
    plt.plot(keys_list, values_list, '--bo')
    plt.suptitle('Liner paragraph distrubition')
    plt.xlabel('entry number of words')
    plt.ylabel('entry number of occurnece of word')
    plt.savefig('statistics.png')


# Task 4
def download_image(url, destination):
    """download the file and save it to a destination file"""
    try:
        urlretrieve(url, destination)
    except Exception as e:
        logger.info("Something went wrong, can't download the image")
        print("Can not download the image")
        exit()


# Task 5
def picture_manipulation():
    """Perform file crop and resize"""
    img_location = "https://www.gutenberg.org/files/64078/64078-h/images/cover.jpg"
    destination = "bookimage.png"
    download_image(img_location, destination)
    # make it bigger
    resize_image(destination, "resizefile.png")
    crop_image(destination, "croppedfile.jpg")
    paste_picture("logo.png", "bookimage.jpg")


# Task 5 a helper
def resize_image(in_file, out_file):
    """resie an image"""
    img = Image.open(in_file)
    w, h = img.size
    new_img = img.resize((int(w/3), int(h/4)))
    new_img.save(out_file)


# Task 5 b helper
def crop_image(in_file, out_file):
    """crop image"""
    img = Image.open(in_file)
    area = (100, 100, 1000, 1000)
    new_img = img.crop(area)
    new_img.save(out_file)


# Task 6
def paste_picture(in_file, out_file):
    """rotate a picture and paste it into
    another one perforemed actions in task 5"""
    logo = Image.open(in_file)
    img_out = Image.open(out_file)
    new_img = img_out.copy()
    location = (30, 30)
    logo.rotate(45)
    new_img.paste(logo, location)
    new_img.save("copy_pasted_" + out_file)


# Task 7
def create_doc():
    """Create a report docx document"""
    title, author = read_authors(File_NAME)
    stat_data = (get_plot_data(File_NAME))
    plot_linear_chart(stat_data)
    doc = docx.Document()
    parent_pargraph1 = doc.add_paragraph("Laboratory 11 Report", "Title")
    parent_pargraph2 = doc.add_paragraph(f"Book Title : {title}\n", "Title")
    doc.add_picture("resizefile.png")
    parent_pargraph3 = doc.add_paragraph(f"Book Author : {author}\n", "Title")
    parent_pargraph3.add_run("\n\nDumyy project makr name")
    parent_pargraph1.italic = True
    parent_pargraph3.bold = True
    current_section = doc.sections[-1]
    current_section.start_type
    parent_pargraph4 = doc.add_paragraph("Ploting page", "Title")
    doc.add_picture("statistics.png")
    parahragraph5 = doc.add_paragraph("The dictionary storing sorted data"
                                      "is as follow\n"
                                      + " ".join(str(e) for e in stat_data))
    parahragraph5.add_run(f"\nThe minimum number of words is {min(stat_data)}"
                          f"\n")
    parahragraph5.add_run(f"\nThe maximum number of words "
                          f"is {max(stat_data)}\n")
    average = (min(stat_data) + max(stat_data)) / 2
    parahragraph5.add_run(f"\nThe average number of words is {average}\n")
    doc.save("report.docx")
    logger.info("Ending the App")


# Main method
def main():
    """Main function"""
    picture_manipulation()
    create_doc()


if __name__ == '__main__':
    main()
